import os
from datetime import datetime
import matplotlib
matplotlib.use('Agg') # GUI 없는 환경에서도 안전하게 차트 생성
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from src.backend.core.database import db_manager

class ReportGenerator:
    def __init__(self, output_dir=r"c:\ameva\AMEVA-STT-Trainer\outputs\reports"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def generate_task_report(self, task_id: str) -> str:
        """태스크 상세 정보, 데이터셋, 로그 등을 종합하여 Word(.docx) 리포트를 생성합니다."""
        task = db_manager.get_task_details(task_id)
        if not task:
            raise ValueError(f"Task ID {task_id} not found in database.")
            
        logs = db_manager.get_logs(task_id=task_id, limit=5000)
        
        doc = Document()
        
        # --- 1. 타이틀 ---
        title = doc.add_heading(f"AMEVA STT Engine - Task Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # --- 2. 기본 정보 ---
        doc.add_heading("1. Task Overview", level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        
        info = [
            ("Task Name", task.get('tsk_nm', 'N/A')),
            ("Task ID", task.get('id', 'N/A')),
            ("Creation Time", task.get('create_dt', 'N/A')),
            ("Status Level", f"Level {task.get('level', 'N/A')} ({task.get('status', 'N/A')})"),
            ("Final Model Path", task.get('model_path', 'Not yet generated'))
        ]
        
        for i, (key, val) in enumerate(info):
            cells = table.rows[i].cells
            cells[0].text = key
            cells[1].text = str(val)
            
        doc.add_paragraph()
            
        # --- 3. 데이터셋 정보 (Chunks) ---
        doc.add_heading("2. Dataset Information", level=1)
        metadatas = task.get('metadatas', [])
        
        if not metadatas:
            doc.add_paragraph("No dataset information found for this task.")
        else:
            for meta in metadatas:
                doc.add_heading(f"Source: {meta.get('file_name', 'Unknown')}", level=2)
                doc.add_paragraph(f"Folder Path: {meta.get('folder_path', 'N/A')}")
                
                chunks = meta.get('chunks', [])
                doc.add_paragraph(f"Total Chunks: {len(chunks)}")
                
                if chunks:
                    chunk_table = doc.add_table(rows=1, cols=3)
                    chunk_table.style = 'Table Grid'
                    hdr_cells = chunk_table.rows[0].cells
                    hdr_cells[0].text = 'Chunk Name'
                    hdr_cells[1].text = 'Path'
                    hdr_cells[2].text = 'Script Snippet'
                    
                    # 10개까지만 표시 (너무 길어짐 방지)
                    for chunk in chunks[:10]:
                        row_cells = chunk_table.add_row().cells
                        row_cells[0].text = chunk.get('chunk_name', '')
                        row_cells[1].text = chunk.get('chunk_path', '')
                        script = chunk.get('script', '')
                        row_cells[2].text = script[:50] + "..." if len(script) > 50 else script
                    
                    if len(chunks) > 10:
                        doc.add_paragraph(f"... and {len(chunks) - 10} more chunks.")
                        
        doc.add_paragraph()

        # --- 4. 학습 메트릭 차트 (Matplotlib) ---
        doc.add_heading("4. Training Metrics & Charts", level=1)
        metrics = db_manager.get_metrics(task_id)
        
        if not metrics:
            doc.add_paragraph("No metric data available for charts.")
        else:
            try:
                import matplotlib.pyplot as plt
                import pandas as pd
                
                df = pd.DataFrame(metrics)
                
                # 1) Loss & Accuracy 차트
                plt.figure(figsize=(8, 4))
                plt.plot(df['step'], df['loss'], label='Loss', color='red', marker='o')
                plt.plot(df['step'], df['accuracy'], label='Accuracy', color='blue', marker='x')
                plt.title('Training Loss & Accuracy over Steps', fontsize=12, fontweight='bold')
                plt.xlabel('Steps')
                plt.ylabel('Value')
                plt.legend()
                plt.grid(True, linestyle='--', alpha=0.6)
                
                chart1_path = os.path.join(self.output_dir, f"chart1_{task_id[:8]}.png")
                plt.savefig(chart1_path, bbox_inches='tight')
                plt.close()
                
                doc.add_picture(chart1_path, width=Inches(6.0))
                doc.add_paragraph("Figure 1: Training Loss and Accuracy trends. Shows how the model converges over time.", style='Caption')
                
                # 2) CPU & Speed 차트
                plt.figure(figsize=(8, 4))
                plt.plot(df['step'], df['cpu_usage'], label='CPU Usage (%)', color='green')
                plt.plot(df['step'], df['speed'], label='Speed (T/s)', color='orange')
                plt.title('System Resource & Processing Speed', fontsize=12, fontweight='bold')
                plt.xlabel('Steps')
                plt.ylabel('Usage / Speed')
                plt.legend()
                plt.grid(True, linestyle='--', alpha=0.6)
                
                chart2_path = os.path.join(self.output_dir, f"chart2_{task_id[:8]}.png")
                plt.savefig(chart2_path, bbox_inches='tight')
                plt.close()
                
                doc.add_picture(chart2_path, width=Inches(6.0))
                doc.add_paragraph("Figure 2: System resources utilized during the training steps.", style='Caption')
                
            except Exception as e:
                doc.add_paragraph(f"[Chart Generation Error] {e}")

        # --- 5. 워크플로우 실행 내역 ---
        doc.add_heading("5. Workflow Details", level=1)
        details = task.get('details', [])
        if not details:
            doc.add_paragraph("No workflow details found.")
        else:
            dtl_table = doc.add_table(rows=1, cols=4)
            dtl_table.style = 'Table Grid'
            hdr_cells = dtl_table.rows[0].cells
            hdr_cells[0].text = 'Step'
            hdr_cells[1].text = 'Name'
            hdr_cells[2].text = 'Status'
            hdr_cells[3].text = 'Parameters'
            
            for dtl in details:
                row_cells = dtl_table.add_row().cells
                row_cells[0].text = str(dtl.get('step_seq', ''))
                row_cells[1].text = dtl.get('step_name', '')
                row_cells[2].text = dtl.get('status', '')
                row_cells[3].text = dtl.get('parameters', '')
                
        doc.add_paragraph()
        
        # --- 6. 생성된 파일 및 폴더 구조 (표 형태) ---
        doc.add_heading("6. Generated Files & Folders", level=1)
        file_table = doc.add_table(rows=1, cols=2)
        file_table.style = 'Table Grid'
        hdr_cells = file_table.rows[0].cells
        hdr_cells[0].text = 'Type'
        hdr_cells[1].text = 'Path / Details'
        
        row_cells = file_table.add_row().cells
        row_cells[0].text = "Model Output"
        row_cells[1].text = task.get('model_path', 'N/A')
        
        doc.add_paragraph()

        # --- 7. 실행 로그 및 예외 (Exceptions) ---
        doc.add_heading("7. Execution Logs & Exceptions", level=1)
        if not logs:
            doc.add_paragraph("No logs found for this task.")
        else:
            log_table = doc.add_table(rows=1, cols=3)
            log_table.style = 'Table Grid'
            hdr_cells = log_table.rows[0].cells
            hdr_cells[0].text = 'Time'
            hdr_cells[1].text = 'Level'
            hdr_cells[2].text = 'Message'
            
            error_count = 0
            for log in logs:
                level = log.get('level', '')
                if level in ['ERROR', 'FAILED', 'CRITICAL']:
                    error_count += 1
                    
                row_cells = log_table.add_row().cells
                row_cells[0].text = log.get('create_dt', '')
                row_cells[1].text = level
                row_cells[2].text = log.get('message', '')
                
            doc.add_paragraph(f"\nTotal Logs: {len(logs)} | Errors/Exceptions: {error_count}")

        # 저장
        report_name = f"Report_{task.get('tsk_nm', 'Task')}_{task_id[:8]}.docx"
        report_name = "".join([c for c in report_name if c.isalpha() or c.isdigit() or c in (' ', '.', '_', '-')]).rstrip()
        save_path = os.path.join(self.output_dir, report_name)
        doc.save(save_path)
        
        # DB에 리포트 위치 업데이트
        db_manager.update_task_status(task_id, level=task.get('level', 1), status=task.get('status', 'RUNNING'), report_path=save_path)
        
        return save_path

report_generator = ReportGenerator()
